import os
import re
from docx import Document

def remove_dash_from_dn(text):
    # Regex to replace 'DN-<digits>' with 'DN<digits>'
    return re.sub(r'DN-(\d+)', r'DN\1', text)

def fix_dn_dash_in_docx(docx_path):
    doc = Document(docx_path)
    modified = False
    
    def process_paragraph(para):
        nonlocal modified
        if 'Debit Note No:' in para.text:
            full_text = ''.join(run.text for run in para.runs)
            new_text = remove_dash_from_dn(full_text)
            if new_text != full_text:
                first_run = para.runs[0] if para.runs else None
                
                # Clear all existing runs
                for run in para.runs:
                    run.text = ''
                
                # Add a new run with the fixed text
                new_run = para.add_run(new_text)
                
                # Copy font formatting from the first run if it exists
                if first_run:
                    new_run.font.bold = first_run.font.bold
                    new_run.font.italic = first_run.font.italic
                    new_run.font.underline = first_run.font.underline
                    new_run.font.color.rgb = first_run.font.color.rgb
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                
                modified = True

    # Process paragraphs in the main document
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    if modified:
        doc.save(docx_path)
        print(f"Fixed {docx_path}")
    else:
        print(f"No changes needed in {docx_path}")

def batch_fix_dn_dash(word_folder):
    docx_files = [f for f in os.listdir(word_folder) if f.lower().endswith('.docx')]
    
    for docx_file in docx_files:
        docx_path = os.path.join(word_folder, docx_file)
        fix_dn_dash_in_docx(docx_path)

if __name__ == "__main__":
    word_folder = r'C:\Users\Avinash Matre\Documents\iLovePDF_Output\ilovepdf-pdf-to-word (2) - Copy'  # update path accordingly
    batch_fix_dn_dash(word_folder)
